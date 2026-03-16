"""
EasyOCR Backend Server
Run with: python ocr_server.py
"""

from fastapi import FastAPI, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
import easyocr
import io
from PIL import Image
import numpy as np
import warnings
import torch
import pdfplumber
# from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from fastapi.responses import FileResponse
import tempfile
import os
import cv2
import fitz  # PyMuPDF
import uuid
import shutil
from spellchecker import SpellChecker
from docx2pdf import convert as convert_to_pdf
try:
    import pythoncom
except ImportError:
    pythoncom = None
import asyncio

# Suppress PyTorch CPU warning
warnings.filterwarnings("ignore", category=UserWarning, module="torch.utils.data.dataloader")

app = FastAPI(title="EasyOCR Server")

# Enable CORS for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def health_check():
    return {"status": "ok", "message": "EasyOCR Server is running"}

# Check for GPU
USE_GPU = torch.cuda.is_available()
print(f"Hardware check: GPU available? {USE_GPU}")
if USE_GPU:
    print(f"Using GPU: {torch.cuda.get_device_name(0)}")
else:
    print("Using CPU (might be slower)")

# Initialize EasyOCR reader (loads model on startup)
print("Loading EasyOCR model...")
reader = easyocr.Reader(['en'], gpu=USE_GPU)
print("EasyOCR model loaded!")

# Initialize SpellChecker
spell = SpellChecker()

def results_to_text_with_spacing(results):
    """
    Convert OCR results to text, preserving vertical spacing by inserting empty lines
    when gaps between lines are detected.
    """
    if not results:
        return ""

    # Sort top-down by Y coordinate
    results.sort(key=lambda x: x[0][0][1])
    
    lines = []
    prev_bottom = -1
    avg_height = 0
    total_height = 0
    count = 0
    
    # Calculate average line height
    for bbox, text, conf in results:
        h = bbox[2][1] - bbox[0][1]
        total_height += h
        count += 1
    
    if count > 0:
        avg_height = total_height / count
    else:
        avg_height = 20 # Default fallback
        
    for bbox, text, conf in results:
        top = bbox[0][1]
        bottom = bbox[2][1]
        
        if prev_bottom != -1:
            gap = top - prev_bottom
            # If gap is significantly larger than line height (indicating paragraph break)
            # We use 0.8 as threshold to catch distinct separations
            if gap > avg_height * 0.8:
                lines.append("") 
                
        lines.append(text)
        prev_bottom = bottom
        
    return "\n".join(lines)

def preprocess_image(image: Image.Image) -> np.ndarray:
    """
    Enhance image for better OCR accuracy.
    """
    # Convert PIL Image to Numpy Array
    img_np = np.array(image)
    
    # Convert to BGR (OpenCV format) if RGB
    if len(img_np.shape) == 3:
        img_np = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)

    # Convert to Grayscale
    gray = cv2.cvtColor(img_np, cv2.COLOR_BGR2GRAY)

    # Upscale if image is small (improves accuracy for small text)
    height, width = gray.shape
    if height < 1000 or width < 1000:
        scale = 2
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    return gray

def run_ocr_on_image(image: Image.Image):
    processed_image = preprocess_image(image)
    
    # Run EasyOCR
    # Optimization: Use fast mode for CPU
    results = reader.readtext(
        processed_image,
        paragraph=False,
        decoder='greedy',
        beamWidth=5,
        batch_size=4 if USE_GPU else 1,
        detail=1
    )
    
    return results

def is_scanned_pdf(pdf_path):
    """Detect if PDF is scanned (images) or digital (selectable text)."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:2]: # Check first 2 pages
                text = page.extract_text()
                if text and len(text.strip()) > 50: # Threshold for "real" text
                    return False
        return True
    except Exception:
        return True # Assume scanned if error

def extract_text_digital_pdf(pdf_path):
    """Extract text from digital PDF directly (fast & clean)."""
    text_content = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_content.append(t)
    return "\n".join(text_content)

@app.get("/health")
def health_check_explicit():
    return {"status": "ok", "message": "EasyOCR Server Running", "gpu_enabled": USE_GPU}

def create_formatted_docx(all_results, output_path, image_width=None):
    """
    Create a Word document from OCR results, attempting to preserve layout using Tab Stops.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # If we know image width, we can map X coordinates to page width (approx 6.5 inches writable)
    # Standard Letter: 8.5in width, 1in margins L/R -> 6.5in writable
    page_width_twips = 6.5 * 1440 # 1 inch = 1440 twips
    
    scale_factor = 1.0
    if image_width and image_width > 0:
        # Scale image width to page width
        scale_factor = (6.5 * 72) / image_width # 72 points per inch approx for coordinate mapping
    
    def get_y(item):
        return item[0][0][1]

    # Sort primarily by Y
    sorted_results = sorted(all_results, key=get_y)
    
    # Group into lines
    lines = []
    current_line = []
    last_y = -1
    y_tolerance = 20  # pixels
    
    for item in sorted_results:
        bbox, text, conf = item
        y = bbox[0][1]
        
        if last_y == -1 or abs(y - last_y) < y_tolerance:
            current_line.append(item)
            if last_y == -1: last_y = y
        else:
            # New line
            current_line.sort(key=lambda x: x[0][0][0])
            lines.append(current_line)
            current_line = [item]
            last_y = y
            
    # Append last line
    if current_line:
        current_line.sort(key=lambda x: x[0][0][0])
        lines.append(current_line)
        
    # Write lines to Docx with Tabs
    for line in lines:
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.tab_stops.clear_all()
        
        # Calculate tab stops
        run_text = ""
        
        for i, (bbox, text, conf) in enumerate(line):
            x_start = bbox[0][0]
            
            # Convert X to Inches for Tab Stop
            # Default assumption if no image_width: 1 pixel = 1/96 inch? 
            # Let's use the scale factor.
            # 1 px * scale_factor (points) / 72 = inches
            
            tab_pos_in = (x_start * scale_factor) / 72.0
            
            # Add tab stop
            try:
                p_format.tab_stops.add_tab_stop(Inches(tab_pos_in), WD_TAB_ALIGNMENT.LEFT)
                run_text += f"\t{text}"
            except:
                # Fallback if calculation fails or negative
                run_text += f" {text}"
            
        p.add_run(run_text)
        
    doc.save(output_path)
    return output_path

import re

