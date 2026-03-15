"""
EasyOCR Backend Server
Run with: python ocr_server.py
"""

from fastapi import FastAPI, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
import easyocr
import io
from PIL import Image
import numpy as np
import warnings
import torch
import pdfplumber
# from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from fastapi.responses import FileResponse
import tempfile
import os
import cv2
import fitz  # PyMuPDF
import uuid
import shutil
from spellchecker import SpellChecker
from docx2pdf import convert as convert_to_pdf
try:
    import pythoncom
except ImportError:
    pythoncom = None
import asyncio

# Suppress PyTorch CPU warning
warnings.filterwarnings("ignore", category=UserWarning, module="torch.utils.data.dataloader")

app = FastAPI(title="EasyOCR Server")

# Enable CORS for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def health_check():
    return {"status": "ok", "message": "EasyOCR Server is running"}

# Check for GPU
USE_GPU = torch.cuda.is_available()
print(f"Hardware check: GPU available? {USE_GPU}")
if USE_GPU:
    print(f"Using GPU: {torch.cuda.get_device_name(0)}")
else:
    print("Using CPU (might be slower)")

# Initialize EasyOCR reader (loads model on startup)
print("Loading EasyOCR model...")
reader = easyocr.Reader(['en'], gpu=USE_GPU)
print("EasyOCR model loaded!")

# Initialize SpellChecker
spell = SpellChecker()

def results_to_text_with_spacing(results):
    """
    Convert OCR results to text, preserving vertical spacing by inserting empty lines
    when gaps between lines are detected.
    """
    if not results:
        return ""

    # Sort top-down by Y coordinate
    results.sort(key=lambda x: x[0][0][1])
    
    lines = []
    prev_bottom = -1
    avg_height = 0
    total_height = 0
    count = 0
    
    # Calculate average line height
    for bbox, text, conf in results:
        h = bbox[2][1] - bbox[0][1]
        total_height += h
        count += 1
    
    if count > 0:
        avg_height = total_height / count
    else:
        avg_height = 20 # Default fallback
        
    for bbox, text, conf in results:
        top = bbox[0][1]
        bottom = bbox[2][1]
        
        if prev_bottom != -1:
            gap = top - prev_bottom
            # If gap is significantly larger than line height (indicating paragraph break)
            # We use 0.8 as threshold to catch distinct separations
            if gap > avg_height * 0.8:
                lines.append("") 
                
        lines.append(text)
        prev_bottom = bottom
        
    return "\n".join(lines)

def preprocess_image(image: Image.Image) -> np.ndarray:
    """
    Enhance image for better OCR accuracy.
    """
    # Convert PIL Image to Numpy Array
    img_np = np.array(image)
    
    # Convert to BGR (OpenCV format) if RGB
    if len(img_np.shape) == 3:
        img_np = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)

    # Convert to Grayscale
    gray = cv2.cvtColor(img_np, cv2.COLOR_BGR2GRAY)

    # Upscale if image is small (improves accuracy for small text)
    height, width = gray.shape
    if height < 1000 or width < 1000:
        scale = 2
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    return gray

def run_ocr_on_image(image: Image.Image):
    processed_image = preprocess_image(image)
    
    # Run EasyOCR
    # Optimization: Use fast mode for CPU
    results = reader.readtext(
        processed_image,
        paragraph=False,
        decoder='greedy',
        beamWidth=5,
        batch_size=4 if USE_GPU else 1,
        detail=1
    )
    
    return results

def is_scanned_pdf(pdf_path):
    """Detect if PDF is scanned (images) or digital (selectable text)."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:2]: # Check first 2 pages
                text = page.extract_text()
                if text and len(text.strip()) > 50: # Threshold for "real" text
                    return False
        return True
    except Exception:
        return True # Assume scanned if error

def extract_text_digital_pdf(pdf_path):
    """Extract text from digital PDF directly (fast & clean)."""
    text_content = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_content.append(t)
    return "\n".join(text_content)

@app.get("/health")
def health_check_explicit():
    return {"status": "ok", "message": "EasyOCR Server Running", "gpu_enabled": USE_GPU}

def create_formatted_docx(all_results, output_path, image_width=None):
    """
    Create a Word document from OCR results, attempting to preserve layout using Tab Stops.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # If we know image width, we can map X coordinates to page width (approx 6.5 inches writable)
    # Standard Letter: 8.5in width, 1in margins L/R -> 6.5in writable
    page_width_twips = 6.5 * 1440 # 1 inch = 1440 twips
    
    scale_factor = 1.0
    if image_width and image_width > 0:
        # Scale image width to page width
        scale_factor = (6.5 * 72) / image_width # 72 points per inch approx for coordinate mapping
    
    def get_y(item):
        return item[0][0][1]

    # Sort primarily by Y
    sorted_results = sorted(all_results, key=get_y)
    
    # Group into lines
    lines = []
    current_line = []
    last_y = -1
    y_tolerance = 20  # pixels
    
    for item in sorted_results:
        bbox, text, conf = item
        y = bbox[0][1]
        
        if last_y == -1 or abs(y - last_y) < y_tolerance:
            current_line.append(item)
            if last_y == -1: last_y = y
        else:
            # New line
            current_line.sort(key=lambda x: x[0][0][0])
            lines.append(current_line)
            current_line = [item]
            last_y = y
            
    # Append last line
    if current_line:
        current_line.sort(key=lambda x: x[0][0][0])
        lines.append(current_line)
        
    # Write lines to Docx with Tabs
    for line in lines:
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.tab_stops.clear_all()
        
        # Calculate tab stops
        run_text = ""
        
        for i, (bbox, text, conf) in enumerate(line):
            x_start = bbox[0][0]
            
            # Convert X to Inches for Tab Stop
            # Default assumption if no image_width: 1 pixel = 1/96 inch? 
            # Let's use the scale factor.
            # 1 px * scale_factor (points) / 72 = inches
            
            tab_pos_in = (x_start * scale_factor) / 72.0
            
            # Add tab stop
            try:
                p_format.tab_stops.add_tab_stop(Inches(tab_pos_in), WD_TAB_ALIGNMENT.LEFT)
                run_text += f"\t{text}"
            except:
                # Fallback if calculation fails or negative
                run_text += f" {text}"
            
        p.add_run(run_text)
        
    doc.save(output_path)
    return output_path

import re
import requests
import json

def is_noise(line):
    """
    Step 1A: Remove garbage lines (Rule-based cleanup)
    """
    # Preserve empty lines (structure)
    if not line.strip():
        return False

    # Remove lines that are too short or have too few letters
    if len(line.strip()) < 4:
        return True
    
    # Count alpha characters
    alpha_count = sum(c.isalpha() for c in line)
    if alpha_count < 3:
        return True
        
    return False

def normalize_line(line):
    """
    Step 1B: Fix spacing & broken punctuation
    """
    # Collapse multiple spaces
    line = re.sub(r'\s+', ' ', line)
    # Fix spacing around punctuation
    line = line.replace(" ,", ",").replace(" .", ".")
    return line.strip()

def correct_spelling(text):
    """
    Step 1C: Rule-based spelling correction
    Preserves newlines but might lose multiple spaces.
    """
    if not text: return ""
    
    lines = text.split('\n')
    corrected_lines = []
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            corrected_lines.append("")
            continue
            
        words = line.split()
        corrected_words = []
        
        for word in words:
            # Strip punctuation for checking
            clean_word = re.sub(r'[^\w\s]', '', word)
            
            # Smart Check:
            # 1. Skip if empty
            # 2. Skip if Capitalized (likely proper noun) BUT check if it's start of sentence?
            #    Simple heuristic: if starts with Uppercase and not 'I', skip.
            # 3. Skip if not in dictionary
            
            if not clean_word:
                corrected_words.append(word)
                continue
                
            # Heuristic: Don't correct proper nouns (Approximate)
            if clean_word[0].isupper() and clean_word.lower() != 'i':
                corrected_words.append(word)
                continue
            
            # Check if word is misspelled
            if len(clean_word) > 1 and clean_word.lower() not in spell:
                # Get the one `most likely` answer
                correction = spell.correction(clean_word)
                
                # If correction found and distinct
                if correction and correction != clean_word.lower():
                    # Preserve original punctuation (simple replacement)
                    # This is tricky: "Word." -> "Correction."
                    corrected_words.append(word.replace(clean_word, correction))
                else:
                    corrected_words.append(word)
            else:
                corrected_words.append(word)
        
        corrected_lines.append(" ".join(corrected_words))
            
    return "\n".join(corrected_lines)

def clean_with_heuristics(text):
    """
    Fallback cleaner:
    1. Merge broken paragraphs (smartly).
    2. Identify headers (Bold).
    3. Correct spelling.
    """
    if not text: return ""
    
    lines = text.split('\n')
    merged_lines = []
    buffer = ""
    
    for line in lines:
        line = line.strip()
        if not line:
            if buffer:
                merged_lines.append(buffer)
                buffer = ""
            merged_lines.append("") # Preserve empty line as paragraph break
            continue
            
        # Header Detection
        is_header = len(line) < 60 and line[0].isupper() and line[-1] not in ".!?,;:"
        
        if is_header:
            if buffer:
                merged_lines.append(buffer)
                buffer = ""
            merged_lines.append(f"**{line}**")
            continue

        # Smart Merging Logic
        should_merge = False
        if buffer:
            # 1. Hyphenation: "com- \n plete" -> "complete"
            if buffer.endswith('-'):
                should_merge = True
                buffer = buffer[:-1] # Remove hyphen
            
            # 2. Sentence Flow:
            # If previous line is "long enough" (likely wrapped) AND doesn't end in punctuation
            # It's probably part of a paragraph.
            elif len(buffer) > 40 and buffer[-1] not in ".!?:":
                should_merge = True
            
            # 3. Lowercase Continuation:
            # If next line starts with lowercase, it belongs to previous sentence.
            elif line and line[0].islower():
                should_merge = True

        if should_merge:
            buffer += " " + line
        else:
            if buffer:
                merged_lines.append(buffer)
            buffer = line
                
    if buffer:
        merged_lines.append(buffer)
        
    # Run spelling on the merged lines
    final_text = "\n".join(merged_lines)
    return correct_spelling(final_text)

def is_ollama_running():
    """Check if Ollama is running and reachable."""
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=1.0)
        return response.status_code == 200
    except:
        return False

def clean_with_ollama(text):
    """
    Step 2: AI cleanup using Ollama
    Requests Markdown output to preserve structure.
    """
    if not text or len(text) < 10:
        return text

    # Quick check to avoid long timeout if Ollama is down
    if not is_ollama_running():
        print("Ollama not detected (or not responding quickly). Skipping AI reconstruction.")
        return clean_with_heuristics(text)

    print("Sending text to Ollama for Layout Reconstruction...")
    
    prompt = """You are an AI document reconstruction expert.
Convert the following OCR text into a clean, structured Markdown document.

Rules:
1. **Layout & Structure**: Preserve the original structure.
2. **Line Breaks**: 
   - **Keep separate lines** for addresses, signatures, lists, and titles. 
   - **Merge lines** ONLY if they form a continuous sentence in a paragraph.
3. **Spacing**: Preserve vertical spacing by using blank lines.
4. **Formatting**: 
   - Use **bold** for headers/titles.
   - Use *italics* for emphasized text.
5. **Tables**: Reconstruct tables using Markdown syntax (| Col1 | Col2 |).
6. **Spelling**: Fix spelling contextually.
   - **IMPORTANT**: Do NOT change proper names (e.g. Piyush Kushwaha).
7. **Output**: Output ONLY Markdown.

Raw Text:
""" + text

    try:
        # Try commonly available models
        # Prioritize installed models: llama3.1
        model = "llama3.1" 
        
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.3 # Lower temperature for more deterministic/faithful output
                }
            },
            timeout=10 # Short timeout: If not super fast (GPU), skip to heuristic for better UX
        )
        
        if response.status_code == 200:
            result = response.json()
            cleaned_text = result.get("response", "").strip()
            print("Ollama layout reconstruction successful!")
            return cleaned_text
        else:
            print(f"Ollama failed with status {response.status_code}. Using heuristic cleaner.")
            return clean_with_heuristics(text)
            
    except Exception as e:
        print(f"Ollama connection failed/slow: {e}. Using heuristic cleaner.")
        return clean_with_heuristics(text)

def clean_and_structure_text(raw_text):
    """
    Clean text, fix line breaks, and structure into paragraphs.
    Implements Rule-based cleanup + AI Layout Reconstruction.
    """
    if not raw_text:
        return ""
        
    lines = raw_text.split('\n')
    
    # 1. Rule-based cleanup (Garbage removal only)
    # We SKIP aggressive merging here to let AI see the original "lines" to detect tables.
    clean_lines = [l for l in lines if not is_noise(l)]
    clean_lines = [normalize_line(l) for l in clean_lines]
    
    # Join with newlines to preserve structure for AI
    combined_text = "\n".join(clean_lines)
    
    # 2. AI Cleanup & Layout Reconstruction
    final_markdown = clean_with_ollama(combined_text)
    
    return final_markdown

def add_markdown_paragraph(doc, text, style='Normal'):
    """
    Add a paragraph to the document with Bold/Italic support parsing.
    Supported syntax: **bold**, *italic*
    """
    p = doc.add_paragraph(style=style)
    # Use tight spacing (0pt) to allow grouping of lines (like addresses).
    # We rely on explicit empty paragraphs for section spacing.
    p.paragraph_format.space_after = Pt(0)
    
    # Split by bold syntax: **text**
    # This regex captures the delimiters and the content
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italics inside non-bold parts (simple nesting not supported)
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*'):
                     run = p.add_run(sub[1:-1])
                     run.italic = True
                else:
                     p.add_run(sub)

def create_docx_from_markdown(markdown_text, output_path):
    """
    Create a Word doc from Markdown text, supporting Headers, Lists, and Tables.
    """
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = markdown_text.split('\n')
    table_buffer = []
    
    def flush_table():
        if not table_buffer: return
        # Create table
        # Filter out separator lines like |---|
        rows = [row for row in table_buffer if not set(row.replace('|','').strip()) <= {'-', ':', ' '}]
        if not rows: return
        
        # Determine max columns
        parsed_rows = [[cell.strip() for cell in row.strip('|').split('|')] for row in rows]
        if not parsed_rows: return
        
        max_cols = max(len(r) for r in parsed_rows)
        if max_cols == 0: return

        table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(parsed_rows):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    row_cells[j].text = cell_text
        
        table_buffer.clear()

    for line in lines:
        stripped = line.strip()
        
        # Table detection
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            continue
        else:
            flush_table()
            
        # Headers
        if stripped.startswith('#'):
            level = len(stripped.split()[0])
            text = stripped.lstrip('#').strip()
            # docx only supports h1-h9
            level = min(level, 9)
            if text:
                h = doc.add_heading(text, level=level)
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(12)
            continue
            
        # Lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            add_markdown_paragraph(doc, stripped[2:], style='List Bullet')
            continue
            
        # Normal text
        if stripped:
            add_markdown_paragraph(doc, stripped)
        else:
            # Add empty paragraph for explicit blank lines in Markdown
            # This helps preserve large gaps
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            
    flush_table()
    doc.save(output_path)
    return output_path

def create_simple_docx(paragraphs, output_path):
    """Create a clean Word doc from paragraphs."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text)
            
    doc.save(output_path)
    return output_path

from pdf2docx import Converter

# Job Storage
JOBS = {}

def update_job(job_id, status, progress=0, result=None, error=None):
    if job_id in JOBS:
        JOBS[job_id].update({
            "status": status,
            "progress": progress,
            "result": result,
            "error": error
        })

def process_document_task(job_id, input_path, filename):
    try:
        update_job(job_id, "Initializing...", 10)
        output_docx = input_path + ".docx"
        
        # 1. Detect PDF Type
        is_pdf = filename.endswith(".pdf")
        
        if is_pdf:
            update_job(job_id, "Analyzing PDF structure...", 20)
            if not is_scanned_pdf(input_path):
                # 2. Digital PDF
                print("Detected Digital PDF.")
                update_job(job_id, "Digital PDF detected. Preserving exact layout...", 30)
                try:
                    cv = Converter(input_path)
                    cv.convert(output_docx)
                    cv.close()
                    update_job(job_id, "Conversion complete!", 100, result=output_docx)
                    return
                except Exception as e:
                    print(f"pdf2docx failed: {e}. Falling back...")
                    update_job(job_id, "Exact match failed. Falling back to AI extraction...", 40)
                    
                    # Fallback
                    raw_text_content = extract_text_digital_pdf(input_path)
                    update_job(job_id, "Reconstructing layout with AI...", 60)
                    markdown_content = clean_and_structure_text(raw_text_content)
                    create_docx_from_markdown(markdown_content, output_docx)
                    update_job(job_id, "Document ready!", 100, result=output_docx)
                    return

            else:
                # 3. Scanned PDF
                print("Detected Scanned PDF.")
                update_job(job_id, "Scanned PDF detected. Starting OCR engine...", 20)
                
                pdf_document = fitz.open(input_path)
                all_page_text = []
                total_pages = len(pdf_document)
                
                for i, page_num in enumerate(range(total_pages)):
                    update_job(job_id, f"Scanning page {i+1} of {total_pages}...", 20 + int((i/total_pages)*40))
                    
                    page = pdf_document.load_page(page_num)
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    results = run_ocr_on_image(img)
                    page_text = results_to_text_with_spacing(results)
                    all_page_text.append(page_text)
                    
                pdf_document.close()
                raw_text_content = "\n\n".join(all_page_text)
                
                # 4. Clean & Structure
                update_job(job_id, "AI is cleaning and reconstructing layout...", 70)
                markdown_content = clean_and_structure_text(raw_text_content)
                
                update_job(job_id, "Generating final Word document...", 90)
                create_docx_from_markdown(markdown_content, output_docx)
                
                update_job(job_id, "Document ready!", 100, result=output_docx)
                return
                
        else:
            # Image file
            update_job(job_id, "Processing Image...", 20)
            image = Image.open(input_path)
            results = run_ocr_on_image(image)
            raw_text_content = results_to_text_with_spacing(results)
            
            update_job(job_id, "AI is cleaning text...", 70)
            markdown_content = clean_and_structure_text(raw_text_content)
            create_docx_from_markdown(markdown_content, output_docx)
            
            update_job(job_id, "Document ready!", 100, result=output_docx)
            return

    except Exception as e:
        import traceback
        traceback.print_exc()
        update_job(job_id, "Failed", 0, error=str(e))

@app.post("/ocr/upload-job")
async def upload_job(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    """Start a background OCR job."""
    job_id = str(uuid.uuid4())
    filename = file.filename.lower() if file.filename else "document"
    
    # Save uploaded file
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_input:
        shutil.copyfileobj(file.file, tmp_input)
        input_path = tmp_input.name
    
    JOBS[job_id] = {
        "status": "Queued",
        "progress": 0,
        "result": None,
        "error": None,
        "filename": filename
    }
    
    background_tasks.add_task(process_document_task, job_id, input_path, filename)
    
    return {"job_id": job_id}

@app.get("/ocr/status/{job_id}")
def get_status(job_id: str):
    if job_id not in JOBS:
        return {"error": "Job not found"}
    return JOBS[job_id]

@app.get("/ocr/download-job/{job_id}")
def download_job(job_id: str):
    if job_id not in JOBS:
        return {"error": "Job not found"}
    
    job = JOBS[job_id]
    if not job["result"]:
        return {"error": "Result not ready"}
        
    return FileResponse(
        job["result"],
        filename=f"{os.path.splitext(job['filename'])[0]}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.post("/ocr/docx")
async def extract_text_docx(file: UploadFile = File(...)):
    """
    Extract text and return a formatted Word document.
    Follows production pipeline:
    1. Detect PDF type (Digital vs Scanned)
    2. Digital -> Convert EXACT layout (pdf2docx) - PRESERVES BOLD, TABLES, SPACES
    3. Scanned -> Convert to Images -> OCR -> Clean & Structure -> Text DOCX
    """
    try:
        contents = await file.read()
        filename = file.filename.lower() if file.filename else "document"
        
        # Create temp file for input
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_input:
            tmp_input.write(contents)
            input_path = tmp_input.name
            
        output_docx = input_path + ".docx"
        
        # 1. Detect PDF Type
        if filename.endswith(".pdf") or contents.startswith(b"%PDF"):
            if not is_scanned_pdf(input_path):
                # 2. Digital PDF: EXACT FORMAT MATCH
                print("Detected Digital PDF. Using pdf2docx for EXACT layout preservation...")
                try:
                    cv = Converter(input_path)
                    # convert(docx_filename, start=0, end=None)
                    cv.convert(output_docx)
                    cv.close()
                    
                    return FileResponse(
                        output_docx, 
                        filename=f"{os.path.splitext(filename)[0]}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    print(f"pdf2docx failed: {e}. Falling back to text extraction.")
                    # Fallback if pdf2docx fails
                    raw_text_content = extract_text_digital_pdf(input_path)
                    print("Cleaning and structuring text (Fallback)...")
                    markdown_content = clean_and_structure_text(raw_text_content)
                    create_docx_from_markdown(markdown_content, output_docx)
                    
            else:
                # 3. Scanned PDF: CLEAN CONTENT (Layout approximation is hard for scans)
                print("Detected Scanned PDF. Running OCR...")
                pdf_document = fitz.open(input_path)
                all_page_text = []
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    # Use 300 DPI for better OCR accuracy as requested
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Run OCR
                    results = run_ocr_on_image(img)
                    
                    # Convert results (boxes) to simple text for this page
                    page_text = results_to_text_with_spacing(results)
                    all_page_text.append(page_text)
                    
                pdf_document.close()
                raw_text_content = "\n\n".join(all_page_text)
                
                # 4. Clean & Structure (Only for Scanned)
                print("Cleaning and structuring text...")
                markdown_content = clean_and_structure_text(raw_text_content)
                create_docx_from_markdown(markdown_content, output_docx)
                
        else:
            # Image file
            print("Processing Image file...")
            image = Image.open(input_path)
            results = run_ocr_on_image(image)
            raw_text_content = results_to_text_with_spacing(results)
            
            # Clean & Structure
            print("Cleaning and structuring text...")
            markdown_content = clean_and_structure_text(raw_text_content)
            create_docx_from_markdown(markdown_content, output_docx)

        return FileResponse(
            output_docx, 
            filename=f"{os.path.splitext(filename)[0]}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

@app.post("/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    """
    Convert Word document (DOCX) to PDF using Microsoft Word (via docx2pdf).
    Requires MS Word to be installed on the server.
    """
    try:
        filename = file.filename
        if not filename:
            filename = "document.docx"
            
        # Create temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_input:
            contents = await file.read()
            tmp_input.write(contents)
            input_path = tmp_input.name
            
        output_path = input_path.replace(".docx", ".pdf")
        
        print(f"Converting Word to PDF: {input_path} -> {output_path}")

        # Run conversion in a separate thread with COM initialization
        def convert_safe():
            pythoncom.CoInitialize()
            try:
                # convert_to_pdf(input_path, output_path) 
                # docx2pdf sometimes fails if output path is explicit? No, it should be fine.
                convert_to_pdf(input_path, output_path)
            except Exception as e:
                print(f"Conversion internal error: {e}")
                raise e
            finally:
                pythoncom.CoUninitialize()

        await asyncio.to_thread(convert_safe)
        
        if not os.path.exists(output_path):
            raise Exception("Conversion failed: Output PDF not created")
            
        return FileResponse(
            output_path,
            filename=f"{os.path.splitext(filename)[0]}.pdf",
            media_type="application/pdf"
        )

    except Exception as e:
        print(f"Word to PDF error: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}

@app.post("/ocr")
async def extract_text(file: UploadFile = File(...)):
    """Extract text from uploaded image or PDF using EasyOCR with optimization"""
    try:
        # Read file content
        contents = await file.read()
        filename = file.filename.lower() if file.filename else ""
        
        all_results = []
        
        # Check if PDF
        if filename.endswith(".pdf") or contents.startswith(b"%PDF"):
            print("Processing PDF file...")
            pdf_document = fitz.open(stream=contents, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                print(f"Processing page {page_num + 1}/{len(pdf_document)}")
                page = pdf_document.load_page(page_num)
                
                # Optimization: Adjust DPI based on hardware
                target_dpi = 200 if not USE_GPU else 300
                pix = page.get_pixmap(dpi=target_dpi) 
                
                # Convert PyMuPDF Pixmap to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Run OCR on this page
                page_results = run_ocr_on_image(img)
                all_results.extend(page_results)
                
            pdf_document.close()
            
        else:
            # Assume Image
            print("Processing Image file...")
            image = Image.open(io.BytesIO(contents))
            all_results = run_ocr_on_image(image)
        
        # Extract text from results
        extracted_lines = []
        total_confidence = 0
        
        for (bbox, text, confidence) in all_results:
            extracted_lines.append(text)
            total_confidence += confidence
        
        full_text = "\n".join(extracted_lines)
        avg_confidence = (total_confidence / len(all_results) * 100) if all_results else 0
        word_count = len(full_text.split())
        
        return {
            "success": True,
            "text": full_text,
            "confidence": round(avg_confidence, 1),
            "wordCount": word_count,
            "lineCount": len(extracted_lines),
            "gpu_used": USE_GPU
        }
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }

if __name__ == "__main__":
    import uvicorn
    print("\n🚀 Starting Optimized EasyOCR Server on http://localhost:8000")
    print("📝 API Endpoint: POST http://localhost:8000/ocr")
    uvicorn.run(app, host="0.0.0.0", port=8000)
